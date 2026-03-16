from io import BytesIO
from datetime import datetime, time, timedelta
from typing import Any

from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.exceptions import PermissionDenied
from rest_framework.response import Response

from apps.matches.models import Match
from apps.tournaments.serializers import MatchSerializer
from apps.tournaments.models import Tournament
from apps.accounts.permissions import IsTournamentCreatorOrAdmin, Role, _get_user_role

from django.db import transaction
from django.db.models import Max

from .models import (
    Schedule,
    ScheduleCourt,
    ScheduleGlobalBreak,
    ScheduleRun,
    ScheduleScope,
    ScheduleScopeCourt,
    ScheduleSlot,
    ScheduleWave,
)
from .serializers import ScheduleSerializer


class ScheduleViewSet(viewsets.ModelViewSet):
    queryset = Schedule.objects.prefetch_related("waves").all()
    serializer_class = ScheduleSerializer
    permission_classes = [AllowAny]

    def get_queryset(self):
        return (
            super()
            .get_queryset()
            .prefetch_related(
                "waves",
                "scopes",
                "scopes__wave",
                "scopes__bound_courts",
                "scopes__bound_courts__court",
            )
            .select_related("created_by")
        )

    def _scoped_tournaments(self, schedule: Schedule):
        return [
            s.tournament
            for s in (
                schedule.scopes.select_related("tournament", "wave")
                .all()
                .order_by("wave__order", "wave_id", "order", "id")
            )
            if s.tournament
        ]

    def _ensure_default_wave_exists(self, schedule: Schedule) -> ScheduleWave:
        w = schedule.waves.all().order_by("order", "id").first()
        if w:
            return w
        return ScheduleWave.objects.create(
            schedule=schedule,
            order=1,
            start_mode=ScheduleWave.StartMode.AFTER_PREVIOUS,
            start_time=None,
            earliest_time=None,
        )

    def _normalize_waves_order(self, schedule: Schedule) -> None:
        waves = list(schedule.waves.all().order_by("order", "id"))
        for idx, w in enumerate(waves, start=1):
            if int(w.order or 0) != idx:
                w.order = idx
                w.save(update_fields=["order"])

    def _ensure_can_use_tournament_in_schedule(self, request, schedule: Schedule, tournament: Tournament) -> None:
        user = getattr(request, "user", None)
        if not user or not getattr(user, "is_authenticated", False):
            raise PermissionDenied("Authentication required")

        role = _get_user_role(user)
        if role == Role.ADMIN or getattr(user, "is_staff", False) or getattr(user, "is_superuser", False):
            return

        if getattr(tournament, "created_by_id", None) != getattr(user, "id", None):
            raise PermissionDenied("You do not have permission to use this tournament in schedule")

    def _validate_schedule_scope_tournament(self, schedule: Schedule, tournament: Tournament) -> None:
        # 1) Дата должна совпадать
        if getattr(schedule, "date", None) and getattr(tournament, "date", None) and schedule.date != tournament.date:
            raise PermissionDenied("Tournament date must match schedule date")

        # 2) Статус должен совпадать правилам draft/official
        if schedule.is_draft:
            if tournament.status != Tournament.Status.CREATED:
                raise PermissionDenied("Draft schedule can include only created tournaments")
        else:
            if tournament.status != Tournament.Status.ACTIVE:
                raise PermissionDenied("Official schedule can include only active tournaments")

        # 3) Все турниры в расписании должны быть одного статуса (created для draft, active для official)
        for t in self._scoped_tournaments(schedule):
            if schedule.is_draft and t.status != Tournament.Status.CREATED:
                raise PermissionDenied("Draft schedule can include only created tournaments")
            if (not schedule.is_draft) and t.status != Tournament.Status.ACTIVE:
                raise PermissionDenied("Official schedule can include only active tournaments")

    def _ensure_can_view_schedule(self, request, schedule: Schedule) -> None:
        tournaments = self._scoped_tournaments(schedule)

        user = getattr(request, "user", None)
        if not user or not getattr(user, "is_authenticated", False):
            for t in tournaments:
                if t.status == Tournament.Status.COMPLETED and t.system == Tournament.System.KING:
                    raise PermissionDenied("Authentication required to view completed King tournaments")

    def _ensure_can_manage_schedule(self, request, schedule: Schedule) -> None:
        user = getattr(request, "user", None)
        if not user or not getattr(user, "is_authenticated", False):
            raise PermissionDenied("Authentication required")

        role = _get_user_role(user)
        if role == Role.ADMIN or getattr(user, "is_staff", False) or getattr(user, "is_superuser", False):
            return

        perm = IsTournamentCreatorOrAdmin()
        tournaments = self._scoped_tournaments(schedule)
        for t in tournaments:
            if not perm.has_object_permission(request, self, t):
                raise PermissionDenied("You do not have permission to manage this schedule")

    def get_permissions(self):
        if self.action in {
            "create",
            "update",
            "partial_update",
            "destroy",
            "save",
            "scopes_add",
            "scopes_remove",
            "scopes_reorder",
            "scopes_update",
            "scopes_move_to_wave",
            "waves_add",
            "waves_update",
            "waves_reorder",
            "waves_remove",
        }:
            return [IsAuthenticated()]
        return super().get_permissions()

    def _validate_bound_courts(self, schedule: Schedule, scope: ScheduleScope, court_ids: list[int]) -> None:
        # Все корты должны принадлежать расписанию
        courts = list(ScheduleCourt.objects.filter(schedule=schedule, id__in=court_ids).values_list("id", flat=True))
        if len(courts) != len(set(court_ids)):
            raise PermissionDenied("Some courts do not belong to this schedule")

        # Корты не должны быть привязаны к другому турниру в этом расписании
        taken = set(
            ScheduleScopeCourt.objects.filter(
                scope__schedule=schedule,
                court_id__in=court_ids,
            )
            .exclude(scope=scope)
            .values_list("court_id", flat=True)
        )
        if taken:
            raise PermissionDenied("Some courts are already bound to another tournament")

    @action(detail=True, methods=["get"], url_path="scopes/available_tournaments", permission_classes=[IsAuthenticated])
    def scopes_available_tournaments(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        q = str(request.query_params.get("q") or "").strip()

        qs = Tournament.objects.all()
        if schedule.date:
            qs = qs.filter(date=schedule.date)

        if schedule.is_draft:
            qs = qs.filter(status=Tournament.Status.CREATED)
        else:
            qs = qs.filter(status=Tournament.Status.ACTIVE)

        # permission filtering
        role = _get_user_role(request.user)
        if role != Role.ADMIN and not getattr(request.user, "is_staff", False) and not getattr(request.user, "is_superuser", False):
            qs = qs.filter(created_by=request.user)

        # exclude already added
        existing_ids = list(ScheduleScope.objects.filter(schedule=schedule).values_list("tournament_id", flat=True))
        if existing_ids:
            qs = qs.exclude(id__in=existing_ids)

        if q:
            qs = qs.filter(name__icontains=q)

        qs = qs.order_by("name", "id")[:200]

        items = [
            {
                "id": int(t.id),
                "name": str(t.name),
                "date": str(t.date) if getattr(t, "date", None) else None,
                "start_time": t.start_time.strftime("%H:%M") if getattr(t, "start_time", None) else None,
                "status": str(t.status),
                "created_by": int(t.created_by_id) if getattr(t, "created_by_id", None) else None,
            }
            for t in qs
        ]

        return Response({"ok": True, "tournaments": items})

    @action(detail=True, methods=["post"], url_path="scopes/add", permission_classes=[IsAuthenticated])
    def scopes_add(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        tournament_id = payload.get("tournament_id")
        if not tournament_id:
            return Response({"ok": False, "error": "no_tournament_id"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            tournament_id_int = int(tournament_id)
        except Exception:
            return Response({"ok": False, "error": "bad_tournament_id"}, status=status.HTTP_400_BAD_REQUEST)

        tournament = Tournament.objects.filter(id=tournament_id_int).first()
        if not tournament:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        self._ensure_can_use_tournament_in_schedule(request, schedule, tournament)
        self._validate_schedule_scope_tournament(schedule, tournament)

        start_mode = str(payload.get("start_mode") or ScheduleScope.StartMode.AFTER_PREVIOUS)
        if start_mode not in {ScheduleScope.StartMode.FIXED, ScheduleScope.StartMode.AFTER_PREVIOUS}:
            start_mode = ScheduleScope.StartMode.AFTER_PREVIOUS

        start_time_value = payload.get("start_time")
        start_time_obj = None
        if start_mode == ScheduleScope.StartMode.FIXED:
            try:
                if isinstance(start_time_value, str):
                    parts = start_time_value.split(":")
                    hh = parts[0]
                    mm = parts[1] if len(parts) > 1 else "0"
                    start_time_obj = time(int(hh), int(mm))
                elif start_time_value is not None:
                    start_time_obj = start_time_value
            except Exception:
                start_time_obj = None
            if start_time_obj is None:
                try:
                    start_time_obj = getattr(tournament, "start_time", None)
                except Exception:
                    start_time_obj = None

        court_ids = payload.get("bound_courts") or []
        try:
            court_ids = [int(x) for x in court_ids]
        except Exception:
            court_ids = []

        with transaction.atomic():
            if ScheduleScope.objects.filter(schedule=schedule, tournament=tournament).exists():
                return Response({"ok": False, "error": "already_added"}, status=status.HTTP_400_BAD_REQUEST)

            default_wave = self._ensure_default_wave_exists(schedule)

            max_order = ScheduleScope.objects.filter(schedule=schedule).aggregate(m=Max("order")).get("m") or 0
            scope = ScheduleScope.objects.create(
                schedule=schedule,
                wave=default_wave,
                tournament=tournament,
                order=int(max_order) + 1,
                start_mode=start_mode,
                start_time=start_time_obj,
            )

            if court_ids:
                self._validate_bound_courts(schedule, scope, court_ids)
                ScheduleScopeCourt.objects.bulk_create(
                    [ScheduleScopeCourt(scope=scope, court_id=c_id) for c_id in sorted(set(court_ids))]
                )

            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="waves/add", permission_classes=[IsAuthenticated])
    def waves_add(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        start_mode = str(payload.get("start_mode") or ScheduleWave.StartMode.AFTER_PREVIOUS)
        if start_mode not in {ScheduleWave.StartMode.FIXED, ScheduleWave.StartMode.AFTER_PREVIOUS}:
            start_mode = ScheduleWave.StartMode.AFTER_PREVIOUS

        def parse_time_value(v):
            if v is None:
                return None
            try:
                if isinstance(v, str):
                    parts = v.split(":")
                    hh = parts[0]
                    mm = parts[1] if len(parts) > 1 else "0"
                    return time(int(hh), int(mm))
                return v
            except Exception:
                return None

        start_time_obj = parse_time_value(payload.get("start_time"))
        earliest_obj = parse_time_value(payload.get("earliest_time"))
        if start_mode == ScheduleWave.StartMode.FIXED and start_time_obj is None:
            return Response({"ok": False, "error": "start_time_required"}, status=status.HTTP_400_BAD_REQUEST)

        max_order = schedule.waves.aggregate(m=Max("order")).get("m") or 0
        ScheduleWave.objects.create(
            schedule=schedule,
            order=int(max_order) + 1,
            start_mode=start_mode,
            start_time=start_time_obj,
            earliest_time=earliest_obj,
        )

        schedule.save(update_fields=["updated_at"])
        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="waves/update", permission_classes=[IsAuthenticated])
    def waves_update(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        wave_id = payload.get("wave_id")
        if not wave_id:
            return Response({"ok": False, "error": "no_wave_id"}, status=status.HTTP_400_BAD_REQUEST)

        wave = ScheduleWave.objects.filter(schedule=schedule, id=wave_id).first()
        if not wave:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        def parse_time_value(v):
            if v is None:
                return None
            try:
                if isinstance(v, str):
                    parts = v.split(":")
                    hh = parts[0]
                    mm = parts[1] if len(parts) > 1 else "0"
                    return time(int(hh), int(mm))
                return v
            except Exception:
                return None

        start_mode = payload.get("start_mode")
        if start_mode is not None:
            start_mode = str(start_mode)
            if start_mode not in {ScheduleWave.StartMode.FIXED, ScheduleWave.StartMode.AFTER_PREVIOUS}:
                return Response({"ok": False, "error": "bad_start_mode"}, status=status.HTTP_400_BAD_REQUEST)
            wave.start_mode = start_mode

        if payload.get("start_time") is not None or wave.start_mode == ScheduleWave.StartMode.FIXED:
            if payload.get("start_time") is not None:
                st = parse_time_value(payload.get("start_time"))
                if st is None:
                    return Response({"ok": False, "error": "bad_start_time"}, status=status.HTTP_400_BAD_REQUEST)
                wave.start_time = st
            if wave.start_mode == ScheduleWave.StartMode.FIXED and wave.start_time is None:
                return Response({"ok": False, "error": "start_time_required"}, status=status.HTTP_400_BAD_REQUEST)
            if wave.start_mode != ScheduleWave.StartMode.FIXED:
                wave.start_time = None

        if payload.get("earliest_time") is not None:
            et = parse_time_value(payload.get("earliest_time"))
            if payload.get("earliest_time") is not None and et is None:
                return Response({"ok": False, "error": "bad_earliest_time"}, status=status.HTTP_400_BAD_REQUEST)
            wave.earliest_time = et

        wave.save(update_fields=["start_mode", "start_time", "earliest_time"])
        schedule.save(update_fields=["updated_at"])
        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="waves/reorder", permission_classes=[IsAuthenticated])
    def waves_reorder(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        wave_ids = payload.get("wave_ids")
        if not isinstance(wave_ids, list) or not wave_ids:
            return Response({"ok": False, "error": "bad_wave_ids"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            wave_ids = [int(x) for x in wave_ids]
        except Exception:
            return Response({"ok": False, "error": "bad_wave_ids"}, status=status.HTTP_400_BAD_REQUEST)

        existing = list(schedule.waves.all().order_by("order", "id"))
        existing_ids = [int(w.id) for w in existing]
        if sorted(existing_ids) != sorted(wave_ids):
            return Response({"ok": False, "error": "wave_ids_mismatch"}, status=status.HTTP_400_BAD_REQUEST)

        with transaction.atomic():
            for idx, wid in enumerate(wave_ids, start=1):
                ScheduleWave.objects.filter(schedule=schedule, id=wid).update(order=idx)
            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="waves/remove", permission_classes=[IsAuthenticated])
    def waves_remove(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        wave_id = payload.get("wave_id")
        if not wave_id:
            return Response({"ok": False, "error": "no_wave_id"}, status=status.HTTP_400_BAD_REQUEST)

        wave = ScheduleWave.objects.filter(schedule=schedule, id=wave_id).first()
        if not wave:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        if schedule.waves.count() <= 1:
            return Response({"ok": False, "error": "cannot_remove_last"}, status=status.HTTP_400_BAD_REQUEST)

        # scopes from removed wave should move to first wave
        target = schedule.waves.exclude(id=wave.id).order_by("order", "id").first()
        if not target:
            return Response({"ok": False, "error": "cannot_remove_last"}, status=status.HTTP_400_BAD_REQUEST)

        with transaction.atomic():
            ScheduleScope.objects.filter(schedule=schedule, wave=wave).update(wave=target)
            wave.delete()
            self._normalize_waves_order(schedule)
            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="scopes/move_to_wave", permission_classes=[IsAuthenticated])
    def scopes_move_to_wave(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        scope_id = payload.get("scope_id")
        wave_id = payload.get("wave_id")
        if not scope_id or not wave_id:
            return Response({"ok": False, "error": "bad_params"}, status=status.HTTP_400_BAD_REQUEST)

        scope = ScheduleScope.objects.filter(schedule=schedule, id=scope_id).select_related("tournament").first()
        if not scope:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        wave = ScheduleWave.objects.filter(schedule=schedule, id=wave_id).first()
        if not wave:
            return Response({"ok": False, "error": "wave_not_found"}, status=status.HTTP_404_NOT_FOUND)

        self._ensure_can_use_tournament_in_schedule(request, schedule, scope.tournament)

        scope.wave = wave
        scope.save(update_fields=["wave"])
        schedule.save(update_fields=["updated_at"])
        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="scopes/remove", permission_classes=[IsAuthenticated])
    def scopes_remove(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        scope_id = payload.get("scope_id")
        tournament_id = payload.get("tournament_id")
        if not scope_id and not tournament_id:
            return Response({"ok": False, "error": "no_scope"}, status=status.HTTP_400_BAD_REQUEST)

        qs = ScheduleScope.objects.filter(schedule=schedule)
        if scope_id:
            qs = qs.filter(id=scope_id)
        else:
            qs = qs.filter(tournament_id=tournament_id)

        scope = qs.select_related("tournament").first()
        if not scope:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        if ScheduleScope.objects.filter(schedule=schedule).count() <= 1:
            return Response({"ok": False, "error": "cannot_remove_last"}, status=status.HTTP_400_BAD_REQUEST)

        self._ensure_can_use_tournament_in_schedule(request, schedule, scope.tournament)

        with transaction.atomic():
            scope.delete()

            # переиндексация order
            scopes = list(ScheduleScope.objects.filter(schedule=schedule).order_by("order", "id"))
            for idx, s in enumerate(scopes, start=1):
                if int(s.order or 0) != idx:
                    s.order = idx
                    s.save(update_fields=["order"])

            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="scopes/reorder", permission_classes=[IsAuthenticated])
    def scopes_reorder(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        order_ids = payload.get("scope_ids")
        if not isinstance(order_ids, list) or not order_ids:
            return Response({"ok": False, "error": "bad_scope_ids"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            order_ids = [int(x) for x in order_ids]
        except Exception:
            return Response({"ok": False, "error": "bad_scope_ids"}, status=status.HTTP_400_BAD_REQUEST)

        existing = list(ScheduleScope.objects.filter(schedule=schedule).order_by("order", "id"))
        existing_ids = [int(s.id) for s in existing]
        if sorted(existing_ids) != sorted(order_ids):
            return Response({"ok": False, "error": "scope_ids_mismatch"}, status=status.HTTP_400_BAD_REQUEST)

        with transaction.atomic():
            for idx, sid in enumerate(order_ids, start=1):
                ScheduleScope.objects.filter(schedule=schedule, id=sid).update(order=idx)
            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="scopes/update", permission_classes=[IsAuthenticated])
    def scopes_update(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        payload = request.data or {}
        scope_id = payload.get("scope_id")
        if not scope_id:
            return Response({"ok": False, "error": "no_scope_id"}, status=status.HTTP_400_BAD_REQUEST)

        scope = ScheduleScope.objects.filter(schedule=schedule, id=scope_id).select_related("tournament").first()
        if not scope:
            return Response({"ok": False, "error": "not_found"}, status=status.HTTP_404_NOT_FOUND)

        self._ensure_can_use_tournament_in_schedule(request, schedule, scope.tournament)
        self._validate_schedule_scope_tournament(schedule, scope.tournament)

        start_mode = payload.get("start_mode")
        start_time_value = payload.get("start_time")

        if start_mode is not None:
            start_mode = str(start_mode)
            if start_mode not in {ScheduleScope.StartMode.FIXED, ScheduleScope.StartMode.AFTER_PREVIOUS}:
                return Response({"ok": False, "error": "bad_start_mode"}, status=status.HTTP_400_BAD_REQUEST)
            scope.start_mode = start_mode

        if scope.start_mode == ScheduleScope.StartMode.FIXED:
            if start_time_value is not None:
                try:
                    if isinstance(start_time_value, str):
                        parts = start_time_value.split(":")
                        hh = parts[0]
                        mm = parts[1] if len(parts) > 1 else "0"
                        scope.start_time = time(int(hh), int(mm))
                    else:
                        scope.start_time = start_time_value
                except Exception:
                    return Response({"ok": False, "error": "bad_start_time"}, status=status.HTTP_400_BAD_REQUEST)
            if scope.start_time is None:
                return Response({"ok": False, "error": "start_time_required"}, status=status.HTTP_400_BAD_REQUEST)
        else:
            scope.start_time = None

        bound_courts = payload.get("bound_courts")
        if bound_courts is not None:
            try:
                court_ids = [int(x) for x in (bound_courts or [])]
            except Exception:
                return Response({"ok": False, "error": "bad_bound_courts"}, status=status.HTTP_400_BAD_REQUEST)

            self._validate_bound_courts(schedule, scope, court_ids)

            with transaction.atomic():
                scope.save(update_fields=["start_mode", "start_time"])
                ScheduleScopeCourt.objects.filter(scope=scope).delete()
                if court_ids:
                    ScheduleScopeCourt.objects.bulk_create(
                        [ScheduleScopeCourt(scope=scope, court_id=c_id) for c_id in sorted(set(court_ids))]
                    )
                schedule.save(update_fields=["updated_at"])

            schedule.refresh_from_db()
            return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

        scope.save(update_fields=["start_mode", "start_time"])
        schedule.save(update_fields=["updated_at"])
        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    def _compute_rr_row_by_team(self, slots_qs: Any) -> dict[tuple[int, int, int], int]:
        rr_row_by_team: dict[tuple[int, int, int], int] = {}
        try:
            from apps.tournaments.models import TournamentEntry

            t_ids: set[int] = set()
            team_ids: set[int] = set()
            for s in slots_qs:
                if getattr(s, "slot_type", None) != "match":
                    continue
                m = getattr(s, "match", None)
                if not m or not getattr(m, "tournament_id", None):
                    continue
                t_ids.add(int(m.tournament_id))
                if getattr(m, "team_1_id", None):
                    team_ids.add(int(m.team_1_id))
                if getattr(m, "team_2_id", None):
                    team_ids.add(int(m.team_2_id))

            if t_ids and team_ids:
                qs = TournamentEntry.objects.filter(tournament_id__in=list(t_ids), team_id__in=list(team_ids)).only(
                    "tournament_id", "team_id", "group_index", "row_index"
                )
                for e in qs:
                    gi = int(e.group_index) if e.group_index is not None else 0
                    if gi <= 0 or e.row_index is None:
                        continue
                    rr_row_by_team[(int(e.tournament_id), gi, int(e.team_id))] = int(e.row_index)
        except Exception:
            rr_row_by_team = {}
        return rr_row_by_team

    def _run_top_label(self, run_obj: Any) -> str:
        try:
            mode = getattr(run_obj, "start_mode", "")
            if mode == "then":
                return "Затем"
            if mode == "not_earlier":
                t = getattr(run_obj, "not_earlier_time", None)
                if t:
                    return f"Не ранее {t.strftime('%H:%M')}"
                return "Не ранее"
            t = getattr(run_obj, "start_time", None)
            if t:
                return t.strftime("%H:%M")
        except Exception:
            pass
        return ""

    def _match_meta_label(self, m: Any, rr_row_by_team: dict[tuple[int, int, int], int]) -> str:
        if not m:
            return ""

        gi = getattr(m, "group_index", None)

        def _is_proam_rr(match_obj: Any) -> bool:
            try:
                t = getattr(match_obj, "tournament", None)
                if not t:
                    return False
                if getattr(t, "system", None) != Tournament.System.ROUND_ROBIN:
                    return False
                name = str(getattr(t, "name", "") or "").lower()
                return ("proam" in name) or ("проам" in name)
            except Exception:
                return False

        def _group_letter(i: Any) -> str:
            try:
                n = int(i)
            except Exception:
                return ""
            letters = [
                "А",
                "Б",
                "В",
                "Г",
                "Д",
                "Е",
                "Ж",
                "З",
                "И",
                "Й",
                "К",
                "Л",
                "М",
                "Н",
                "О",
                "П",
                "Р",
                "С",
                "Т",
                "У",
                "Ф",
                "Х",
                "Ц",
                "Ч",
                "Ш",
                "Щ",
                "Э",
                "Ю",
                "Я",
            ]
            if n <= 0:
                return ""
            if n <= len(letters):
                return letters[n - 1]
            return str(n)

        if gi is None:
            group = ""
        else:
            group = f"гр.{_group_letter(gi)}" if _is_proam_rr(m) else f"гр.{gi}"

        try:
            system = getattr(getattr(m, "tournament", None), "system", "")
        except Exception:
            system = ""

        if system == "knockout":
            rn = str(getattr(m, "round_name", "") or "").strip()
            if rn:
                return rn
            ri = getattr(m, "round_index", None)
            return f"Раунд {ri}".strip() if ri is not None else ""

        if system == "round_robin":
            rn = str(getattr(m, "round_name", "") or "").strip()
            if rn and "гр." in rn and "•" in rn:
                return rn
            try:
                gi_int = int(gi) if gi is not None else 0
                t_id = int(getattr(m, "tournament_id", 0) or 0)
                a_id = int(getattr(m, "team_1_id", 0) or 0)
                b_id = int(getattr(m, "team_2_id", 0) or 0)
                r1 = rr_row_by_team.get((t_id, gi_int, a_id))
                r2 = rr_row_by_team.get((t_id, gi_int, b_id))
                pair = f"{r1}-{r2}" if (r1 is not None and r2 is not None) else ""
                return " • ".join([p for p in [group, pair] if p])
            except Exception:
                return group

        if system == "king":
            def _raw_team_label(team: Any) -> str:
                if not team:
                    return ""
                return str(
                    getattr(team, "display_name", None)
                    or getattr(team, "full_name", None)
                    or getattr(team, "name", None)
                    or ""
                ).strip()

            def _normalize_pair(text: str) -> str:
                return str(text or "").replace(" ", "").replace("/", "+").upper()

            a = _normalize_pair(_raw_team_label(getattr(m, "team_1", None)))
            b = _normalize_pair(_raw_team_label(getattr(m, "team_2", None)))
            vs = f"{a} vs {b}" if a and b else ""
            rn = str(getattr(m, "round_name", "") or "").strip()
            mid = getattr(m, "id", None)
            fallback = f"Матч {mid}" if mid is not None else ""
            meta = vs or rn or fallback
            return " • ".join([p for p in [group, meta] if p])

        return group

    def destroy(self, request, *args, **kwargs):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)
        return super().destroy(request, *args, **kwargs)

    @action(detail=True, methods=["post"], url_path="save", permission_classes=[IsAuthenticated])
    def save(self, request, pk=None):
        """Сохранить расписание одним запросом.

        MVP-реализация: полная перезапись дочерних сущностей (courts/runs/slots/global_breaks).

        Payload:
        - match_duration_minutes?: int
        - courts: [{ index:int, name:str, first_start_time?:"HH:MM"|null }]
        - runs: [{ index:int, start_mode:"fixed"|"then"|"not_earlier", start_time?:"HH:MM"|null, not_earlier_time?:"HH:MM"|null }]
        - slots: [{ run_index:int, court_index:int, slot_type:"match"|"text", match_id?:int|null,
                   text_title?:str|null, text_subtitle?:str|null, override_title?:str|null, override_subtitle?:str|null }]
        - global_breaks: [{ position:int, time:"HH:MM", text:str }]
        """

        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)
        payload = request.data or {}

        def parse_time(v):
            if v is None or v == "":
                return None
            if isinstance(v, time):
                return v
            if isinstance(v, str):
                parts = v.split(":")
                if len(parts) >= 2:
                    hh = parts[0]
                    mm = parts[1]
                    ss = parts[2] if len(parts) >= 3 else "0"
                    try:
                        return time(int(hh), int(mm), int(ss))
                    except Exception:
                        return time(int(hh), int(mm))
                return None
            return None

        courts_in = payload.get("courts") or []
        runs_in = payload.get("runs") or []
        slots_in = payload.get("slots") or []
        breaks_in = payload.get("global_breaks") or []

        try:
            mdm = payload.get("match_duration_minutes")
            if mdm is not None:
                schedule.match_duration_minutes = int(mdm)
        except Exception:
            return Response(
                {"ok": False, "error": "bad_params", "detail": "match_duration_minutes должен быть числом"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        from .models import ScheduleCourt, ScheduleGlobalBreak, ScheduleRun, ScheduleSlot

        with transaction.atomic():
            schedule.save(update_fields=["match_duration_minutes", "updated_at"])

            # Preserve scope->court bindings by court index.
            # `save` recreates courts, so court IDs change and bindings would be lost otherwise.
            from .models import ScheduleScopeCourt

            preserved_bindings: list[tuple[int, int]] = list(
                ScheduleScopeCourt.objects.filter(scope__schedule=schedule)
                .select_related("court")
                .values_list("scope_id", "court__index")
            )

            # очищаем дочерние записи
            ScheduleSlot.objects.filter(schedule=schedule).delete()
            ScheduleGlobalBreak.objects.filter(schedule=schedule).delete()
            ScheduleRun.objects.filter(schedule=schedule).delete()
            ScheduleCourt.objects.filter(schedule=schedule).delete()

            # создаём корты
            court_by_index = {}
            for item in courts_in:
                idx = int(item.get("index"))
                name = item.get("name") or f"Корт {idx}"
                fst = parse_time(item.get("first_start_time"))
                court = ScheduleCourt.objects.create(
                    schedule=schedule,
                    index=idx,
                    name=str(name)[:128],
                    first_start_time=fst,
                )
                court_by_index[idx] = court

            # Restore preserved bindings on recreated courts.
            if preserved_bindings and court_by_index:
                restored: list[ScheduleScopeCourt] = []
                for scope_id, court_index in preserved_bindings:
                    court_obj = court_by_index.get(int(court_index))
                    if not court_obj:
                        continue
                    restored.append(ScheduleScopeCourt(scope_id=int(scope_id), court_id=int(court_obj.id)))
                if restored:
                    ScheduleScopeCourt.objects.bulk_create(restored, ignore_conflicts=True)

            # создаём запуски
            run_by_index = {}
            for item in runs_in:
                idx = int(item.get("index"))
                start_mode = item.get("start_mode") or "fixed"
                st = parse_time(item.get("start_time"))
                ne = parse_time(item.get("not_earlier_time"))
                run = ScheduleRun.objects.create(
                    schedule=schedule,
                    index=idx,
                    start_mode=start_mode,
                    start_time=st,
                    not_earlier_time=ne,
                )
                run_by_index[idx] = run

            # глобальные паузы
            for item in breaks_in:
                pos = int(item.get("position"))
                t = parse_time(item.get("time"))
                txt = str(item.get("text") or "")[:256]
                if not t or not txt:
                    continue
                ScheduleGlobalBreak.objects.create(schedule=schedule, position=pos, time=t, text=txt)

            # слоты
            for item in slots_in:
                run_idx = int(item.get("run_index"))
                court_idx = int(item.get("court_index"))
                run = run_by_index.get(run_idx)
                court = court_by_index.get(court_idx)
                if not run or not court:
                    continue
                slot_type = item.get("slot_type")
                match_id = item.get("match_id")
                if match_id is not None:
                    try:
                        match_id = int(match_id)
                    except Exception:
                        match_id = None

                ScheduleSlot.objects.create(
                    schedule=schedule,
                    run=run,
                    court=court,
                    slot_type=slot_type,
                    match_id=match_id,
                    text_title=item.get("text_title"),
                    text_subtitle=item.get("text_subtitle"),
                    override_title=item.get("override_title"),
                    override_subtitle=item.get("override_subtitle"),
                )

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="runs/add", permission_classes=[IsAuthenticated])
    def add_run(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        from .models import ScheduleRun

        with transaction.atomic():
            last = ScheduleRun.objects.filter(schedule=schedule).order_by("-index").first()
            next_index = int(getattr(last, "index", 0) or 0) + 1
            ScheduleRun.objects.create(
                schedule=schedule,
                index=next_index,
                start_mode=ScheduleRun.StartMode.THEN,
                start_time=None,
                not_earlier_time=None,
            )
            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["post"], url_path="runs/delete", permission_classes=[IsAuthenticated])
    def delete_run(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        run_id = (request.data or {}).get("run_id")
        try:
            run_id = int(run_id)
        except Exception:
            return Response({"ok": False, "error": "bad_params", "detail": "run_id обязателен"}, status=status.HTTP_400_BAD_REQUEST)

        from .models import ScheduleRun, ScheduleSlot

        run = ScheduleRun.objects.filter(schedule=schedule, id=run_id).first()
        if not run:
            return Response({"ok": False, "error": "not_found", "detail": "Запуск не найден"}, status=status.HTTP_404_NOT_FOUND)

        last = ScheduleRun.objects.filter(schedule=schedule).order_by("-index").first()
        if last and int(last.index) != int(run.index):
            return Response(
                {"ok": False, "error": "not_last", "detail": "Можно удалить только последний запуск"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Run is considered non-empty if ANY slot in this run has a match or any text/override content.
        run_slots = ScheduleSlot.objects.filter(schedule=schedule, run=run)
        has_content = run_slots.exclude(match_id__isnull=True).exists() or run_slots.exclude(
            text_title__isnull=True,
            text_subtitle__isnull=True,
            override_title__isnull=True,
            override_subtitle__isnull=True,
        ).exclude(
            text_title="",
            text_subtitle="",
            override_title="",
            override_subtitle="",
        ).exists()

        if has_content:
            return Response(
                {"ok": False, "error": "not_empty", "detail": "Нельзя удалить запуск: в нём есть данные"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        with transaction.atomic():
            # delete slots first (CASCADE from run would also work, but be explicit)
            run_slots.delete()
            run.delete()
            schedule.save(update_fields=["updated_at"])

        schedule.refresh_from_db()
        return Response({"ok": True, "schedule": ScheduleSerializer(schedule).data})

    @action(detail=True, methods=["get"], url_path="matches_pool", permission_classes=[AllowAny])
    def matches_pool(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_view_schedule(request, schedule)

        tournament_ids = list(schedule.scopes.values_list("tournament_id", flat=True))
        if not tournament_ids:
            return Response({"ok": True, "matches": []})

        assigned_ids = set(
            schedule.slots.exclude(match_id__isnull=True).values_list("match_id", flat=True)
        )

        # Пул матчей:
        # - все НЕ завершенные матчи турниров/стадий из scope (кроме placement)
        # - плюс завершенные матчи, если они уже назначены в расписание (для таймлайна/просмотра счёта)
        base_qs = Match.objects.filter(tournament_id__in=tournament_ids).exclude(stage=Match.Stage.PLACEMENT)
        pool_qs = (
            base_qs.exclude(status=Match.Status.COMPLETED)
            | base_qs.filter(id__in=assigned_ids, status=Match.Status.COMPLETED)
        )
        pool_qs = (
            pool_qs.select_related("tournament", "team_1", "team_2", "winner")
            .prefetch_related("sets")
            .order_by("tournament_id", "stage", "round_index", "order_in_round", "id")
        )

        results: list[dict[str, Any]] = []
        for m in pool_qs:
            dto = MatchSerializer(m).data
            dto["is_assigned"] = m.id in assigned_ids
            results.append(dto)

        return Response({"ok": True, "matches": results})

    @action(detail=True, methods=["get"], url_path="live_state", permission_classes=[AllowAny])
    def live_state(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_view_schedule(request, schedule)
        match_ids = list(
            schedule.slots.exclude(match_id__isnull=True).values_list("match_id", flat=True).distinct()
        )
        if not match_ids:
            return Response({"ok": True, "matches": []})

        qs = (
            Match.objects.filter(id__in=match_ids)
            .prefetch_related("sets")
            .only("id", "status", "started_at", "finished_at")
        )
        items = []
        for m in qs:
            sets = []
            try:
                for s in m.sets.all():
                    sets.append(
                        {
                            "games_1": s.games_1,
                            "games_2": s.games_2,
                            "tb_1": s.tb_1,
                            "tb_2": s.tb_2,
                            "is_tiebreak_only": getattr(s, "is_tiebreak_only", False),
                        }
                    )
            except Exception:
                sets = []
            items.append(
                {
                    "id": m.id,
                    "status": m.status,
                    "started_at": m.started_at,
                    "finished_at": m.finished_at,
                    "sets": sets,
                }
            )
        return Response({"ok": True, "matches": items})

    @action(detail=True, methods=["get"], url_path="export/pdf", permission_classes=[IsAuthenticated])
    def export_pdf(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        # Prefer HTML/CSS -> PDF rendering via headless Chromium (maximally identical to browser).
        # If Playwright/Chromium is unavailable on the server, fall back to legacy ReportLab renderer.
        try:
            from playwright.sync_api import sync_playwright

            courts = list(schedule.courts.all().order_by("index"))
            dense_mode = len(courts) >= 9
            runs = list(schedule.runs.all().order_by("index"))
            slots_qs = schedule.slots.select_related(
                "run",
                "court",
                "match",
                "match__tournament",
                "match__team_1",
                "match__team_2",
            )
            slots_map: dict[tuple[int, int], Any] = {}
            for s in slots_qs:
                slots_map[(s.run_id, s.court_id)] = s

            rr_row_by_team = self._compute_rr_row_by_team(slots_qs)

            tournament_names = []
            for scope in schedule.scopes.select_related("tournament").all():
                if scope.tournament and scope.tournament.name:
                    tournament_names.append(str(scope.tournament.name))
            title = " + ".join(tournament_names) if tournament_names else "Расписание"

            def esc(v: Any) -> str:
                import html

                return html.escape("" if v is None else str(v))

            def _raw_team_label(team: Any) -> str:
                if not team:
                    return ""
                return (
                    getattr(team, "display_name", None)
                    or getattr(team, "name", None)
                    or str(team)
                )

            def _split_players(label: str) -> list[tuple[str, str]]:
                import re

                s = (label or "").strip()
                if not s or s.upper() == "TBD":
                    return []

                parts = [p.strip() for p in re.split(r"\s*/\s*", s) if p.strip()]
                out: list[tuple[str, str]] = []
                for p in parts:
                    tokens = [t for t in p.split() if t]
                    if not tokens:
                        continue
                    surname = tokens[0]
                    initial = ""
                    if len(tokens) >= 2 and tokens[1]:
                        initial = tokens[1][0].upper()
                    out.append((surname, initial))
                return out

            def _players_unique_key(surname: str, initial: str) -> str:
                return f"{surname}|{initial}" if initial else surname

            surname_to_keys: dict[str, set[str]] = {}
            for s in slots_qs:
                if getattr(s, "slot_type", None) != "match":
                    continue
                m = getattr(s, "match", None)
                if not m:
                    continue
                for team in [getattr(m, "team_1", None), getattr(m, "team_2", None)]:
                    raw = _raw_team_label(team)
                    for (sn, ini) in _split_players(raw):
                        surname_to_keys.setdefault(sn, set()).add(_players_unique_key(sn, ini))

            def team_label(team: Any) -> str:
                if not team:
                    return "TBD"

                raw = _raw_team_label(team)
                if not raw:
                    return "TBD"

                players = _split_players(raw)
                if not players:
                    return raw

                labels: list[str] = []
                for (sn, ini) in players:
                    if len(surname_to_keys.get(sn, set())) > 1 and ini:
                        labels.append(f"{sn} {ini}.")
                    else:
                        labels.append(sn)
                return "\n".join(labels)

            def run_top_label(run_obj: Any) -> str:
                return self._run_top_label(run_obj)

            def match_meta_label(m: Any) -> str:
                return self._match_meta_label(m, rr_row_by_team)

            def slot_class(slot: Any) -> str:
                try:
                    if slot and slot.match and slot.match.status == Match.Status.LIVE:
                        return " cellLive"
                    if slot and slot.match and slot.match.status == Match.Status.COMPLETED:
                        return " cellCompleted"
                except Exception:
                    pass
                return ""

            def slot_html(run_obj: Any, slot: Any) -> str:
                if not slot:
                    return ""
                if getattr(slot, "slot_type", None) == "text":
                    return f'<div class="cellText">{esc(getattr(slot, "text_title", "") or "")}</div>'
                if getattr(slot, "slot_type", None) == "match" and getattr(slot, "match_id", None):
                    if getattr(slot, "override_title", None):
                        return f'<div class="cellText">{esc(slot.override_title)}</div>'
                    m = getattr(slot, "match", None)
                    if not m:
                        return f'<div class="cellText">Матч #{esc(slot.match_id)}</div>'
                    t0 = str(getattr(slot, "override_subtitle", None) or "").strip() or run_top_label(run_obj)
                    meta = match_meta_label(m)
                    t1 = team_label(getattr(m, "team_1", None))
                    t2 = team_label(getattr(m, "team_2", None))
                    t1_lines = [ln for ln in str(t1).split("\n") if ln]
                    t2_lines = [ln for ln in str(t2).split("\n") if ln]
                    t1_html = "".join([f'<div class="player">{esc(ln)}</div>' for ln in t1_lines])
                    t2_html = "".join([f'<div class="player">{esc(ln)}</div>' for ln in t2_lines])
                    return (
                        '<div class="matchTile">'
                        + (f'<div class="matchTop">{esc(t0)}</div>' if t0 else '')
                        + (f'<div class="matchMeta">{esc(meta)}</div>' if meta else '')
                        + f'<div class="teamBlock">{t1_html}</div>'
                        '<div class="vs">против</div>'
                        f'<div class="teamBlock">{t2_html}</div>'
                        "</div>"
                    )
                return ""

            court_headers = "".join(
                [
                    "<th>"
                    + (
                        f"<div class=\"courtName\">{esc((c.name or '').replace('Корт ', 'Корт') or f'Корт{c.index}')}</div>"
                        if dense_mode
                        else f"<div class=\"courtName\">{esc(c.name or f'Корт {c.index}')}</div>"
                    )
                    + "</th>"
                    for c in courts
                ]
            )

            body_rows: list[str] = []
            for r in runs:
                left = (
                    f"<td class=\"runCol\">"
                    f"<div class=\"runTitle\">Запуск {esc(r.index)}</div>"
                    + (
                        f"<div class=\"runPlan\">План: {esc(r.start_time.strftime('%H:%M'))}</div>"
                        if getattr(r, "start_time", None)
                        else ""
                    )
                    + "</td>"
                )
                cells = []
                for c in courts:
                    slot = slots_map.get((r.id, c.id))
                    cells.append(f"<td class=\"cell{slot_class(slot)}\">{slot_html(r, slot)}</td>")
                body_rows.append("<tr>" + left + "".join(cells) + "</tr>")

            html_doc = f"""<!doctype html>
<html lang=\"ru\">
<head>
  <meta charset=\"utf-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
  <title>{esc(title)}</title>
  <style>
    @page {{ size: A4 portrait; margin: 10mm; }}
    * {{ box-sizing: border-box; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, 'Noto Sans', 'DejaVu Sans', sans-serif;
      color: #111827;
    }}

    body.dense {{ font-size: 11px; }}
    body.dense thead th {{ padding: 4px 4px; }}
    body.dense tbody td {{ padding: 4px; }}
    .pageHeader {{ display:flex; align-items:flex-start; justify-content:space-between; margin-bottom: 6px; }}
    .pageHeader .h1 {{ font-size: 16px; font-weight: 700; line-height: 1.15; }}
    .pageHeader .h2 {{ font-size: 11px; color: #374151; margin-top: 2px; }}
    .pageHeader .date {{ font-size: 11px; color: #111827; }}

    table {{ width: 100%; border-collapse: collapse; table-layout: {'auto' if dense_mode else 'fixed'}; }}
    thead th {{
      background: #F3F4F6;
      border-bottom: 1px solid #D1D5DB;
      padding: 6px 6px;
      text-align: left;
      vertical-align: bottom;
    }}
    thead th:first-child {{ width: {'72px' if dense_mode else '88px'}; }}
    tbody td {{
      border-bottom: 1px solid #E5E7EB;
      border-right: 1px solid #E5E7EB;
      padding: 6px;
      vertical-align: middle;
    }}
    tbody td:last-child {{ border-right: 0; }}
    tbody tr td:first-child {{ border-right: 1px solid #D1D5DB; }}

    .courtName {{ font-weight: 700; }}
    .runCol .runTitle {{ font-weight: 700; }}
    .runCol .runPlan {{ font-size: 10px; color: #6B7280; margin-top: 1px; }}

    .cell {{ text-align: center; }}
    .matchTile {{ display:flex; flex-direction:column; align-items:center; justify-content:flex-start; }}
    .matchTop {{ font-size: 10px; color: #6B7280; margin-bottom: 2px; }}
    .matchMeta {{ font-size: 10px; color: #374151; margin-bottom: 2px; width: 100%; text-align: right; }}
    .teamBlock {{ display:flex; flex-direction:column; align-items:center; justify-content:flex-start; }}
    .player {{ font-size: {'11px' if dense_mode else '12px'}; line-height: 1.12; }}
    .vs {{ font-weight: 700; font-size: {'10px' if dense_mode else '11px'}; margin: 2px 0; }}
    .cellText {{ font-size: {'11px' if dense_mode else '12px'}; line-height: 1.15; }}

    .cellLive {{ background: #D1FAE5; }}
    .cellCompleted {{ background: #E5E7EB; }}
  </style>
</head>
<body class="{'dense' if dense_mode else ''}">
  <div class="pageHeader">
    <div>
      <div class="h1">Расписание</div>
      <div class="h2">{esc(title)}</div>
    </div>
    <div class=\"date\">{esc(schedule.date)}</div>
  </div>

  <table>
    <thead>
      <tr>
        <th>Запуск</th>
        {court_headers}
      </tr>
    </thead>
    <tbody>
      {''.join(body_rows)}
    </tbody>
  </table>
</body>
</html>"""

            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page(viewport={"width": 900, "height": 1400})
                page.set_content(html_doc, wait_until="load")
                page.emulate_media(media="screen")
                pdf_bytes = page.pdf(format="A4", landscape=False, print_background=True)
                browser.close()

            from django.http import HttpResponse

            resp = HttpResponse(pdf_bytes, content_type="application/pdf")
            resp["Content-Disposition"] = f'attachment; filename="schedule_{schedule.id}.pdf"'
            return resp
        except Exception:
            pass

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfgen.canvas import Canvas
        except Exception:
            return Response(
                {
                    "ok": False,
                    "error": "pdf_export_unavailable",
                    "detail": "PDF export is unavailable on this server (missing dependency: reportlab or playwright)",
                },
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        # Подготовим данные (все runs/courts + mapping слотов)
        courts = list(schedule.courts.all().order_by("index"))
        runs = list(schedule.runs.all().order_by("index"))
        slots_qs = schedule.slots.select_related("run", "court", "match", "match__tournament", "match__team_1", "match__team_2")
        slots_map: dict[tuple[int, int], Any] = {}
        for s in slots_qs:
            slots_map[(s.run_id, s.court_id)] = s

        breaks = list(schedule.global_breaks.all().order_by("position"))
        breaks_by_pos: dict[int, list[Any]] = {}
        for br in breaks:
            breaks_by_pos.setdefault(int(br.position), []).append(br)

        tournament_names = []
        for scope in schedule.scopes.select_related("tournament").all():
            if scope.tournament and scope.tournament.name:
                tournament_names.append(str(scope.tournament.name))
        title = " + ".join(tournament_names) if tournament_names else "Расписание"

        rr_row_by_team = self._compute_rr_row_by_team(slots_qs)

        def run_top_label(run_obj: Any) -> str:
            return self._run_top_label(run_obj)

        def match_meta_label(m: Any) -> str:
            return self._match_meta_label(m, rr_row_by_team)

        # --- PDF canvas ---
        buf = BytesIO()
        page_size = landscape(A4)
        c = Canvas(buf, pagesize=page_size)

        # Регистрируем шрифт с кириллицей. Если не найден, используем стандартный.
        font_name = "Helvetica"
        try:
            import os

            candidates = [
                ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                ("Arial", r"C:\\Windows\\Fonts\\arial.ttf"),
                ("Arial", r"C:\\Windows\\Fonts\\ARIAL.TTF"),
            ]

            for name, path in candidates:
                if path and os.path.exists(path):
                    pdfmetrics.registerFont(TTFont(name, path))
                    font_name = name
                    break
        except Exception:
            pass

        page_w, page_h = page_size
        margin_x = 24
        margin_y = 24

        header_h = 48
        col_header_h = 34
        run_row_h = 110
        break_row_h = 24

        # Доступная высота под таблицу (без шапки)
        table_top_y = page_h - margin_y - header_h
        table_bottom_y = margin_y
        table_h = table_top_y - table_bottom_y

        # Ширины колонок
        run_col_w = 110

        def split_courts_evenly(items: list[Any], max_per_page: int = 10) -> list[list[Any]]:
            n = len(items)
            if n <= max_per_page:
                return [items]
            # Равномерное разбиение на k страниц
            import math

            k = int(math.ceil(n / max_per_page))
            per = int(math.ceil(n / k))
            chunks = []
            i = 0
            while i < n:
                chunks.append(items[i : i + per])
                i += per
            return chunks

        court_chunks = split_courts_evenly(courts, max_per_page=10)

        def draw_page_header():
            c.setFillColor(colors.black)
            c.setFont(font_name, 14)
            c.drawString(margin_x, page_h - margin_y - 18, "Расписание")
            c.setFont(font_name, 11)
            c.drawString(margin_x, page_h - margin_y - 36, f"{title}")
            c.drawRightString(page_w - margin_x, page_h - margin_y - 18, str(schedule.date))

        def draw_table_header(courts_subset: list[Any], x0: float, y0: float, col_w: float):
            # фон
            c.setFillColor(colors.HexColor("#F3F4F6"))
            c.rect(x0, y0 - col_header_h, run_col_w + col_w * len(courts_subset), col_header_h, fill=1, stroke=0)

            c.setFillColor(colors.black)
            c.setFont(font_name, 10)
            c.drawString(x0 + 6, y0 - 18, "Запуск")
            for idx, court in enumerate(courts_subset):
                cx = x0 + run_col_w + idx * col_w
                name = court.name or f"Корт {court.index}"
                c.drawString(cx + 6, y0 - 16, name[:24])
                try:
                    if getattr(court, "first_start_time", None):
                        c.setFont(font_name, 9)
                        c.setFillColor(colors.HexColor("#4B5563"))
                        c.drawString(cx + 6, y0 - 29, f"Начало {court.first_start_time.strftime('%H:%M')}")
                        c.setFillColor(colors.black)
                        c.setFont(font_name, 10)
                except Exception:
                    pass

            c.setStrokeColor(colors.HexColor("#D1D5DB"))
            c.line(x0, y0 - col_header_h, x0 + run_col_w + col_w * len(courts_subset), y0 - col_header_h)

        def slot_text(run_obj: Any, slot: Any) -> str:
            if not slot:
                return ""
            if slot.slot_type == "text":
                return (slot.text_title or "")[:40]
            if slot.slot_type == "match" and slot.match_id:
                # основной текст: override_title или display team names
                if slot.override_title:
                    return str(slot.override_title)[:40]
                m = slot.match
                if not m:
                    return f"Матч #{slot.match_id}"
                top = str(getattr(slot, "override_subtitle", None) or "").strip() or run_top_label(run_obj)
                meta = match_meta_label(m)
                t1 = getattr(m.team_1, "display_name", None) or getattr(m.team_1, "name", None) or str(m.team_1) if m.team_1 else "TBD"
                t2 = getattr(m.team_2, "display_name", None) or getattr(m.team_2, "name", None) or str(m.team_2) if m.team_2 else "TBD"
                lines: list[str] = []
                if top:
                    lines.append(top)
                if meta:
                    lines.append(meta)
                lines.extend([str(t1), "против", str(t2)])
                return "\n".join(lines)
            return ""

        def slot_status_color(slot: Any):
            try:
                if slot and slot.match and slot.match.status == Match.Status.LIVE:
                    return colors.HexColor("#D1FAE5")
                if slot and slot.match and slot.match.status == Match.Status.COMPLETED:
                    return colors.HexColor("#E5E7EB")
            except Exception:
                pass
            return None

        def draw_run_row(courts_subset: list[Any], run_obj: Any, x0: float, y_top: float, col_w: float):
            # левая колонка
            c.setFont(font_name, 10)
            c.setFillColor(colors.black)
            label = f"Запуск {run_obj.index}"
            time_str = ""
            if run_obj.start_time:
                time_str = run_obj.start_time.strftime("%H:%M")
            c.drawString(x0 + 6, y_top - 18, (label)[:24])
            if time_str:
                c.setFont(font_name, 9)
                c.setFillColor(colors.HexColor("#4B5563"))
                c.drawString(x0 + 6, y_top - 34, f"План: {time_str}")
                c.setFillColor(colors.black)
                c.setFont(font_name, 10)

            # ячейки
            for idx, court in enumerate(courts_subset):
                cx = x0 + run_col_w + idx * col_w
                slot = slots_map.get((run_obj.id, court.id))

                fill = slot_status_color(slot)
                if fill is not None:
                    c.setFillColor(fill)
                    c.rect(cx, y_top - run_row_h, col_w, run_row_h, fill=1, stroke=0)

                c.setFillColor(colors.black)
                txt = slot_text(run_obj, slot)
                if txt:
                    lines = [ln.strip() for ln in str(txt).split("\n") if ln.strip()]
                    # вертикальное размещение ближе к центру ячейки
                    c.setFont(font_name, 10)
                    start_y = y_top - 42
                    line_h = 14
                    for i, ln in enumerate(lines[:5]):
                        # грубое центрирование по ширине
                        w = pdfmetrics.stringWidth(ln, font_name, 10)
                        tx = cx + max(4, (col_w - w) / 2.0)
                        c.drawString(tx, start_y - i * line_h, ln[:60])

            # линии сетки
            c.setStrokeColor(colors.HexColor("#E5E7EB"))
            c.line(x0, y_top - run_row_h, x0 + run_col_w + col_w * len(courts_subset), y_top - run_row_h)

        def draw_break_row(courts_subset: list[Any], br: Any, x0: float, y_top: float, col_w: float):
            w = run_col_w + col_w * len(courts_subset)
            c.setFillColor(colors.HexColor("#DBEAFE"))
            c.rect(x0, y_top - break_row_h, w, break_row_h, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#1E3A8A"))
            c.setFont(font_name, 10)
            t = br.time.strftime("%H:%M") if br.time else ""
            c.drawString(x0 + 6, y_top - 17, f"{t} — {br.text}"[:80])
            c.setStrokeColor(colors.HexColor("#93C5FD"))
            c.line(x0, y_top - break_row_h, x0 + w, y_top - break_row_h)

        # Рендер страниц
        for courts_subset in court_chunks:
            # ширина под корты
            usable_w = page_w - 2 * margin_x
            col_w = max(60.0, (usable_w - run_col_w) / max(1, len(courts_subset)))

            # сколько строк помещается по высоте
            # (грубая оценка: шапка колонок + строки запусков + строки пауз)
            y = table_top_y
            draw_page_header()
            draw_table_header(courts_subset, margin_x, y, col_w)
            y -= col_header_h

            # Печатаем с разбиением по высоте
            # В position 0 допускаем паузу "перед первым запуском"
            run_index_to_obj = {r.index: r for r in runs}

            current_run_idx = 1
            while current_run_idx <= len(runs):
                # Паузы перед запуском current_run_idx
                for br in breaks_by_pos.get(current_run_idx - 1, []):
                    if y - break_row_h < table_bottom_y:
                        c.showPage()
                        draw_page_header()
                        y = table_top_y
                        draw_table_header(courts_subset, margin_x, y, col_w)
                        y -= col_header_h
                    draw_break_row(courts_subset, br, margin_x, y, col_w)
                    y -= break_row_h

                run_obj = run_index_to_obj.get(current_run_idx)
                if run_obj is None:
                    current_run_idx += 1
                    continue

                if y - run_row_h < table_bottom_y:
                    c.showPage()
                    draw_page_header()
                    y = table_top_y
                    draw_table_header(courts_subset, margin_x, y, col_w)
                    y -= col_header_h

                draw_run_row(courts_subset, run_obj, margin_x, y, col_w)
                y -= run_row_h
                current_run_idx += 1

            # Паузы после последнего запуска — position == len(runs)
            for br in breaks_by_pos.get(len(runs), []):
                if y - break_row_h < table_bottom_y:
                    c.showPage()
                    draw_page_header()
                    y = table_top_y
                    draw_table_header(courts_subset, margin_x, y, col_w)
                    y -= col_header_h
                draw_break_row(courts_subset, br, margin_x, y, col_w)
                y -= break_row_h

            c.showPage()

        c.save()
        pdf = buf.getvalue()
        buf.close()

        from django.http import HttpResponse

        resp = HttpResponse(pdf, content_type="application/pdf")
        resp["Content-Disposition"] = f'attachment; filename="schedule_{schedule.id}.pdf"'
        return resp

    @action(detail=True, methods=["get"], url_path="export/docx", permission_classes=[IsAuthenticated])
    def export_docx(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except Exception:
            return Response(
                {
                    "ok": False,
                    "error": "docx_export_unavailable",
                    "detail": "DOCX export is unavailable on this server (missing dependency: python-docx)",
                },
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        courts = list(schedule.courts.all().order_by("index"))
        runs = list(schedule.runs.all().order_by("index"))
        slots_qs = schedule.slots.select_related(
            "run",
            "court",
            "match",
            "match__tournament",
            "match__team_1",
            "match__team_2",
        )
        slots_map: dict[tuple[int, int], Any] = {}
        for s in slots_qs:
            slots_map[(s.run_id, s.court_id)] = s

        breaks = list(schedule.global_breaks.all().order_by("position"))
        breaks_by_pos: dict[int, list[Any]] = {}
        for br in breaks:
            breaks_by_pos.setdefault(int(br.position), []).append(br)

        tournament_names = []
        for scope in schedule.scopes.select_related("tournament").all():
            if scope.tournament and scope.tournament.name:
                tournament_names.append(str(scope.tournament.name))
        title = " + ".join(tournament_names) if tournament_names else "Расписание"

        rr_row_by_team = self._compute_rr_row_by_team(slots_qs)

        def run_top_label(run_obj: Any) -> str:
            return self._run_top_label(run_obj)

        def match_meta_label(m: Any) -> str:
            return self._match_meta_label(m, rr_row_by_team)

        def split_courts_evenly(items: list[Any], max_per_page: int = 10) -> list[list[Any]]:
            n = len(items)
            if n <= max_per_page:
                return [items]
            import math

            k = int(math.ceil(n / max_per_page))
            per = int(math.ceil(n / k))
            chunks = []
            i = 0
            while i < n:
                chunks.append(items[i : i + per])
                i += per
            return chunks

        def _raw_team_label(team: Any) -> str:
            if not team:
                return ""
            return getattr(team, "display_name", None) or getattr(team, "name", None) or str(team)

        def _split_players(label: str) -> list[tuple[str, str]]:
            import re

            s = (label or "").strip()
            if not s or s.upper() == "TBD":
                return []

            parts = [p.strip() for p in re.split(r"\s*/\s*", s) if p.strip()]
            out: list[tuple[str, str]] = []
            for p in parts:
                tokens = [t for t in p.split() if t]
                if not tokens:
                    continue
                surname = tokens[0]
                initial = ""
                if len(tokens) >= 2 and tokens[1]:
                    initial = tokens[1][0].upper()
                out.append((surname, initial))
            return out

        def _players_unique_key(surname: str, initial: str) -> str:
            return f"{surname}|{initial}" if initial else surname

        surname_to_keys: dict[str, set[str]] = {}
        for s in slots_qs:
            if getattr(s, "slot_type", None) != "match":
                continue
            m = getattr(s, "match", None)
            if not m:
                continue
            for team in [getattr(m, "team_1", None), getattr(m, "team_2", None)]:
                raw = _raw_team_label(team)
                for (sn, ini) in _split_players(raw):
                    surname_to_keys.setdefault(sn, set()).add(_players_unique_key(sn, ini))

        def team_label(team: Any) -> list[str]:
            if not team:
                return ["TBD"]

            raw = _raw_team_label(team)
            if not raw:
                return ["TBD"]

            players = _split_players(raw)
            if not players:
                return [raw]

            labels: list[str] = []
            for (sn, ini) in players:
                if len(surname_to_keys.get(sn, set())) > 1 and ini:
                    labels.append(f"{sn} {ini}.")
                else:
                    labels.append(sn)
            return labels

        def slot_cell_lines(run_obj: Any, slot: Any) -> list[str]:
            if not slot:
                return []
            if getattr(slot, "slot_type", None) == "text":
                t = str(getattr(slot, "override_title", None) or getattr(slot, "text_title", None) or "").strip()
                return [t] if t else []
            if getattr(slot, "slot_type", None) == "match" and getattr(slot, "match_id", None):
                if getattr(slot, "override_title", None):
                    return [str(slot.override_title)]
                m = getattr(slot, "match", None)
                if not m:
                    return [f"Матч #{slot.match_id}"]
                top = str(getattr(slot, "override_subtitle", None) or "").strip() or run_top_label(run_obj)
                meta = match_meta_label(m)
                a = team_label(getattr(m, "team_1", None))
                b = team_label(getattr(m, "team_2", None))
                lines: list[str] = []
                if top:
                    lines.append(top)
                if meta:
                    lines.append(meta)
                lines.extend([ln for ln in a if ln])
                lines.append("против")
                lines.extend([ln for ln in b if ln])
                return lines
            return []

        court_chunks = split_courts_evenly(courts, max_per_page=10)

        doc = Document()

        # A4 landscape like PDF (default in current PDF fallback)
        try:
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
        except Exception:
            pass

        def add_header_block():
            p = doc.add_paragraph("Расписание")
            try:
                p.runs[0].font.size = Pt(14)
                p.runs[0].bold = True
            except Exception:
                pass
            p2 = doc.add_paragraph(title)
            try:
                p2.runs[0].font.size = Pt(10)
            except Exception:
                pass
            p3 = doc.add_paragraph(str(schedule.date))
            try:
                p3.runs[0].font.size = Pt(10)
            except Exception:
                pass

        for chunk_i, courts_subset in enumerate(court_chunks):
            if chunk_i > 0:
                try:
                    doc.add_page_break()
                except Exception:
                    doc.add_paragraph("\f")

            add_header_block()

            # table: header row + body (runs + breaks)
            cols = 1 + len(courts_subset)
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Запуск"
            for j, cobj in enumerate(courts_subset, start=1):
                hdr_cells[j].text = str(getattr(cobj, "name", None) or f"Корт {getattr(cobj, 'index', j)}")
                try:
                    hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                except Exception:
                    pass

            def add_run_row(run_obj: Any):
                row_cells = table.add_row().cells
                row_cells[0].text = f"Запуск {run_obj.index}"
                # (План) show start_time if present
                if getattr(run_obj, "start_time", None):
                    row_cells[0].text += f"\nПлан: {run_obj.start_time.strftime('%H:%M')}"

                for j, cobj in enumerate(courts_subset, start=1):
                    slot = slots_map.get((run_obj.id, cobj.id))
                    lines = slot_cell_lines(run_obj, slot)
                    row_cells[j].text = "\n".join([str(x) for x in lines if str(x).strip()]) if lines else ""
                    try:
                        row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for par in row_cells[j].paragraphs:
                            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in par.runs:
                                r.font.size = Pt(10)
                    except Exception:
                        pass

            def add_break_row(br: Any):
                row_cells = table.add_row().cells
                t = br.time.strftime("%H:%M") if getattr(br, "time", None) else ""
                row_cells[0].text = ""
                msg = f"{t} — {br.text}".strip(" —")
                if cols >= 2:
                    row_cells[1].text = msg
                    # best-effort merge across remaining court columns
                    try:
                        for k in range(2, cols):
                            row_cells[1].merge(row_cells[k])
                    except Exception:
                        pass
                    try:
                        for par in row_cells[1].paragraphs:
                            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for r in par.runs:
                                r.font.size = Pt(10)
                    except Exception:
                        pass

            run_index_to_obj = {int(r.index): r for r in runs}
            current_run_idx = 1
            while current_run_idx <= len(runs):
                for br in breaks_by_pos.get(current_run_idx - 1, []):
                    add_break_row(br)
                run_obj = run_index_to_obj.get(current_run_idx)
                if run_obj is not None:
                    add_run_row(run_obj)
                current_run_idx += 1

            for br in breaks_by_pos.get(len(runs), []):
                add_break_row(br)

            doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        from django.http import HttpResponse

        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="schedule_{schedule.id}.docx"'
        return resp

    @action(detail=True, methods=["get"], url_path="export/xlsx", permission_classes=[IsAuthenticated])
    def export_xlsx(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font, PatternFill
        except Exception:
            return Response(
                {
                    "ok": False,
                    "error": "xlsx_export_unavailable",
                    "detail": "XLSX export is unavailable on this server (missing dependency: openpyxl)",
                },
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        courts = list(schedule.courts.all().order_by("index"))
        runs = list(schedule.runs.all().order_by("index"))
        slots_qs = schedule.slots.select_related(
            "run",
            "court",
            "match",
            "match__tournament",
            "match__team_1",
            "match__team_2",
        )
        slots_map: dict[tuple[int, int], Any] = {}
        for s in slots_qs:
            slots_map[(s.run_id, s.court_id)] = s

        breaks = list(schedule.global_breaks.all().order_by("position"))
        breaks_by_pos: dict[int, list[Any]] = {}
        for br in breaks:
            breaks_by_pos.setdefault(int(br.position), []).append(br)

        rr_row_by_team = self._compute_rr_row_by_team(slots_qs)

        tournament_names = []
        for scope in schedule.scopes.select_related("tournament").all():
            if scope.tournament and scope.tournament.name:
                tournament_names.append(str(scope.tournament.name))
        title = " + ".join(tournament_names) if tournament_names else "Расписание"

        def run_top_label(run_obj: Any) -> str:
            return self._run_top_label(run_obj)

        def match_meta_label(m: Any) -> str:
            return self._match_meta_label(m, rr_row_by_team)

        def slot_cell_text(run_obj: Any, slot: Any) -> str:
            if not slot:
                return ""
            if getattr(slot, "slot_type", None) == "text":
                return str(getattr(slot, "override_title", None) or getattr(slot, "text_title", None) or "").strip()
            if getattr(slot, "slot_type", None) == "match" and getattr(slot, "match_id", None):
                if getattr(slot, "override_title", None):
                    return str(slot.override_title)
                m = getattr(slot, "match", None)
                if not m:
                    return f"Матч #{getattr(slot, 'match_id', '')}".strip()
                top = str(getattr(slot, "override_subtitle", None) or "").strip() or run_top_label(run_obj)
                meta = match_meta_label(m)
                t1 = (
                    getattr(getattr(m, "team_1", None), "display_name", None)
                    or getattr(getattr(m, "team_1", None), "name", None)
                    or str(getattr(m, "team_1", None) or "TBD")
                )
                t2 = (
                    getattr(getattr(m, "team_2", None), "display_name", None)
                    or getattr(getattr(m, "team_2", None), "name", None)
                    or str(getattr(m, "team_2", None) or "TBD")
                )
                parts = [p for p in [top, meta, t1, "против", t2] if str(p).strip()]
                return "\n".join([str(p) for p in parts])
            return ""

        wb = Workbook()
        ws = wb.active
        ws.title = "Расписание"

        header_fill = PatternFill("solid", fgColor="F3F4F6")
        live_fill = PatternFill("solid", fgColor="D1FAE5")
        completed_fill = PatternFill("solid", fgColor="E5E7EB")
        break_fill = PatternFill("solid", fgColor="DBEAFE")

        ws["A1"] = "Расписание"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A2"] = title
        ws["A3"] = str(schedule.date)

        row = 5
        ws.cell(row=row, column=1, value="Запуск").font = Font(bold=True)
        ws.cell(row=row, column=1).fill = header_fill
        ws.cell(row=row, column=1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        for i, cobj in enumerate(courts, start=2):
            name = str(getattr(cobj, "name", None) or f"Корт {getattr(cobj, 'index', i-1)}")
            cell = ws.cell(row=row, column=i, value=name)
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        row += 1
        run_index_to_obj = {int(r.index): r for r in runs}
        current_run_idx = 1
        while current_run_idx <= len(runs):
            for br in breaks_by_pos.get(current_run_idx - 1, []):
                t = br.time.strftime("%H:%M") if getattr(br, "time", None) else ""
                msg = f"{t} — {br.text}".strip(" —")
                ws.cell(row=row, column=1, value="")
                if courts:
                    ws.cell(row=row, column=2, value=msg)
                    if len(courts) >= 2:
                        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=1 + len(courts))
                ws.cell(row=row, column=2).fill = break_fill
                ws.cell(row=row, column=2).alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                row += 1

            run_obj = run_index_to_obj.get(current_run_idx)
            if run_obj is None:
                current_run_idx += 1
                continue

            left = f"Запуск {run_obj.index}"
            if getattr(run_obj, "start_time", None):
                left += f"\nПлан: {run_obj.start_time.strftime('%H:%M')}"
            ws.cell(row=row, column=1, value=left).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

            for i, cobj in enumerate(courts, start=2):
                slot = slots_map.get((run_obj.id, cobj.id))
                txt = slot_cell_text(run_obj, slot)
                c = ws.cell(row=row, column=i, value=txt)
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

                try:
                    if slot and slot.match and slot.match.status == Match.Status.LIVE:
                        c.fill = live_fill
                    elif slot and slot.match and slot.match.status == Match.Status.COMPLETED:
                        c.fill = completed_fill
                except Exception:
                    pass

            row += 1
            current_run_idx += 1

        for br in breaks_by_pos.get(len(runs), []):
            t = br.time.strftime("%H:%M") if getattr(br, "time", None) else ""
            msg = f"{t} — {br.text}".strip(" —")
            ws.cell(row=row, column=1, value="")
            if courts:
                ws.cell(row=row, column=2, value=msg)
                if len(courts) >= 2:
                    ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=1 + len(courts))
            ws.cell(row=row, column=2).fill = break_fill
            ws.cell(row=row, column=2).alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            row += 1

        # Column widths (best-effort)
        ws.column_dimensions["A"].width = 18
        for i in range(2, 2 + max(1, len(courts))):
            from openpyxl.utils import get_column_letter

            ws.column_dimensions[get_column_letter(i)].width = 26

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)

        from django.http import HttpResponse

        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        resp["Content-Disposition"] = f'attachment; filename="schedule_{schedule.id}.xlsx"'
        return resp

    @action(detail=True, methods=["get"], url_path="export/numbers", permission_classes=[IsAuthenticated])
    def export_numbers(self, request, pk=None):
        """Экспорт для Numbers (iPad/macOS): отдаём XLSX, который Numbers открывает и редактирует."""
        # Переиспользуем XLSX-экспорт, но с другим именем файла.
        schedule: Schedule = self.get_object()
        self._ensure_can_manage_schedule(request, schedule)
        resp = self.export_xlsx(request, pk=pk)
        try:
            resp["Content-Disposition"] = f'attachment; filename="schedule_{schedule.id}_numbers.xlsx"'
        except Exception:
            pass
        return resp

    @action(detail=True, methods=["get"], url_path="planned_times", permission_classes=[AllowAny])
    def planned_times(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_view_schedule(request, schedule)

        duration = int(getattr(schedule, "match_duration_minutes", 40) or 40)
        runs = list(schedule.runs.all().order_by("index"))
        if not runs:
            return Response({"ok": True, "runs": []})

        def base_start_time() -> time:
            court = schedule.courts.order_by("index").first()
            if court and court.first_start_time:
                return court.first_start_time
            first = runs[0]
            if first.start_time:
                return first.start_time
            return time(10, 0)

        t0 = base_start_time()
        current_dt = datetime.combine(schedule.date, t0)

        items = []
        for r in runs:
            if r.start_mode == "fixed" and r.start_time:
                current_dt = datetime.combine(schedule.date, r.start_time)
            elif r.start_mode == "not_earlier" and r.not_earlier_time:
                candidate = datetime.combine(schedule.date, r.not_earlier_time)
                if candidate > current_dt:
                    current_dt = candidate
            # then: используем текущее current_dt

            items.append(
                {
                    "index": r.index,
                    "planned_start_time": current_dt.time().strftime("%H:%M"),
                }
            )

            current_dt = current_dt + timedelta(minutes=duration)

        return Response({"ok": True, "runs": items, "match_duration_minutes": duration})

    @action(detail=True, methods=["get"], url_path="conflicts", permission_classes=[AllowAny])
    def conflicts(self, request, pk=None):
        schedule: Schedule = self.get_object()
        self._ensure_can_view_schedule(request, schedule)

        slots = (
            schedule.slots.exclude(match_id__isnull=True)
            .select_related("run", "court", "match", "match__team_1", "match__team_2")
            .order_by("run__index", "court__index")
        )

        by_run: dict[int, list[Any]] = {}
        for s in slots:
            by_run.setdefault(int(s.run.index), []).append(s)

        result_runs = []
        for run_index, run_slots in by_run.items():
            players_map: dict[int, list[dict[str, Any]]] = {}
            for s in run_slots:
                m = s.match
                if not m:
                    continue

                player_ids = []
                t1 = getattr(m, "team_1", None)
                t2 = getattr(m, "team_2", None)
                if t1:
                    if getattr(t1, "player_1_id", None):
                        player_ids.append(int(t1.player_1_id))
                    if getattr(t1, "player_2_id", None):
                        player_ids.append(int(t1.player_2_id))
                if t2:
                    if getattr(t2, "player_1_id", None):
                        player_ids.append(int(t2.player_1_id))
                    if getattr(t2, "player_2_id", None):
                        player_ids.append(int(t2.player_2_id))

                for pid in [p for p in player_ids if p]:
                    players_map.setdefault(pid, []).append(
                        {
                            "court_index": int(s.court.index),
                            "match_id": int(m.id),
                            "slot_id": int(s.id),
                        }
                    )

            conflicts_players = []
            for pid, occ in players_map.items():
                if len(occ) > 1:
                    conflicts_players.append({"player_id": pid, "occurrences": occ})

            if conflicts_players:
                result_runs.append({"run_index": run_index, "players": conflicts_players})

        return Response({"ok": True, "runs": result_runs})
